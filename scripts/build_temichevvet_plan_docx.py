#!/usr/bin/env python3
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PLAN_PATH = ROOT / "docs/product/temichevvet-improvement-plan.json"
OUTPUT_DIR = ROOT / "output/docx"

NAVY = "173A5E"
TEAL = "1F7A83"
BLUE = "2E74B5"
INK = "18324A"
MUTED = "637789"
LIGHT = "EEF4F7"
LIGHT_TEAL = "E7F3F3"
GRID = "CBD8E2"
WHITE = "FFFFFF"
GOLD = "946200"
RED = "9B1C1C"

STATUS_COLORS = {
    "PLANNED": ("E9EEF3", INK),
    "IN_PROGRESS": ("FFF0C2", GOLD),
    "CODE_READY": ("DDEBFA", BLUE),
    "LOCAL_VERIFIED": ("DDF2F1", TEAL),
    "CLINIC_ACCEPTED": ("DFF1E2", "27643B"),
    "STABLE": ("CFE8D5", "1F5A32"),
    "BLOCKED": ("F8DEDE", RED),
}


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run(run, *, size=10.5, color=INK, bold=False, italic=False, font="Arial"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def shade(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), color)


def cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = tc_mar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, *, size=9.2, color=INK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    set_run(p.add_run(text), size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_margins(cell)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def keep_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def set_repeat_heading(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


def add_page_field(paragraph):
    for kind, text in (("begin", None), (None, "PAGE"), ("separate", None), (None, "1"), ("end", None)):
        if kind:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), kind)
            paragraph.add_run()._r.append(node)
        elif text == "PAGE":
            node = OxmlElement("w:instrText")
            node.set(qn("xml:space"), "preserve")
            node.text = " PAGE "
            paragraph.add_run()._r.append(node)
        else:
            set_run(paragraph.add_run(text), size=8, color=MUTED)


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # LibreOffice may carry a hanging indent across a page break when the
    # paragraph has no explicit indentation.  Keep headings slightly inside
    # the text area so the first character is never clipped in DOCX/PDF.
    p.paragraph_format.left_indent = Cm(0.12)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 11)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    p.paragraph_format.keep_with_next = True
    sizes = {1: 14, 2: 13, 3: 11.5}
    colors = {1: NAVY, 2: BLUE, 3: TEAL}
    set_run(p.add_run(text), size=sizes[level], color=colors[level], bold=True)
    set_repeat_heading(p)
    return p


def compact_summary_action(text: str, limit: int = 150) -> str:
    normalized = " ".join(text.split())
    if len(normalized) <= limit:
        return normalized
    candidate = normalized[: limit + 1]
    for marker in (". ", "; ", ": "):
        split_at = candidate.rfind(marker)
        if split_at >= 80:
            return candidate[: split_at + 1].rstrip() + " …"
    split_at = candidate.rfind(" ")
    return candidate[:split_at].rstrip() + " …"


def add_para(
    doc,
    text: str,
    *,
    size=10.5,
    color=INK,
    bold=False,
    italic=False,
    after=6,
    align=None,
    line_spacing=1.2,
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    set_run(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text: str, *, level=0, color=INK):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.65 + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.18
    set_run(p.add_run("• "), size=10.2, color=TEAL, bold=True)
    set_run(p.add_run(text), size=10.2, color=color)
    return p


def add_status_badge(doc, code: str, label: str):
    fill, color = STATUS_COLORS[code]
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(5.0)
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_text(cell, label.upper(), size=8.8, color=color, bold=True)
    cell_margins(cell, top=65, bottom=65)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)


def add_callout(doc, title: str, body: str, fill=LIGHT_TEAL, accent=TEAL, trailing=True):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(title), size=10.5, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    set_run(p2.add_run(body), size=9.8, color=INK)
    cell_margins(cell, top=140, start=180, bottom=140, end=180)
    if trailing:
        doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_item_detail(doc, item, status_labels, *, page_break=False):
    if page_break:
        doc.add_page_break()
    add_heading(doc, f"{item['id']}. {item['title']}", 1)
    add_status_badge(doc, item["status"], status_labels[item["status"]])
    add_heading(doc, "Что есть сейчас", 2)
    add_para(doc, item["current_state"])
    add_heading(doc, "Результат для пользователя", 2)
    add_para(doc, item["outcome"])
    add_heading(doc, "Когда считаем готовым", 2)
    for criterion in item["acceptance_criteria"]:
        add_bullet(doc, criterion)
    add_heading(doc, "Следующий шаг", 2)
    add_callout(doc, "Ближайшее действие", item["next_action"])
    if item.get("dependencies"):
        add_para(doc, "Зависимости: " + "; ".join(item["dependencies"]) + ".", size=9.3, color=MUTED)
    if item.get("evidence"):
        add_para(
            doc,
            "Зафиксированные факты: " + " ".join(item["evidence"]),
            size=8.0,
            color=MUTED,
            italic=True,
            after=1,
            line_spacing=1.05,
        )


def configure_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(p.add_run("TEMICHEVVET  ·  ПРОДУКТОВАЯ ДОРОЖНАЯ КАРТА"), size=8, color=MUTED, bold=True)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), GRID)
    borders.append(bottom)
    p_pr.append(borders)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(fp.add_run("План доработки TemichevVet  ·  "), size=8, color=MUTED)
    add_page_field(fp)


def build():
    plan = json.loads(PLAN_PATH.read_text(encoding="utf-8"))
    status_labels = {item["code"]: item["label"] for item in plan["status_model"]}
    items = {item["id"]: item for item in plan["items"]}

    doc = Document()
    configure_document(doc)

    add_para(doc, "ЖИВАЯ ДОРОЖНАЯ КАРТА", size=10, color=TEAL, bold=True, after=3)
    add_para(doc, plan["title"], size=29, color=NAVY, bold=True, after=6)
    add_para(doc, "От функционально широкой CRM — к принятому и устойчивому клиническому продукту", size=13.2, color=MUTED, after=18)

    meta = doc.add_table(rows=2, cols=2)
    meta.autofit = False
    values = [
        ("Актуально", plan["updated_at"]),
        ("Версия", plan["version"]),
        ("Проверенный срез", plan["source_snapshot"]),
        ("Команда статуса", "make plan"),
    ]
    for index, (label, value) in enumerate(values):
        cell = meta.cell(index // 2, index % 2)
        shade(cell, LIGHT)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        set_run(p.add_run(label.upper()), size=7.8, color=MUTED, bold=True)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        set_run(p2.add_run(value), size=10.2, color=INK, bold=True)
        cell_margins(cell, top=120, start=150, bottom=120, end=150)

    add_heading(doc, "Цель", 1)
    add_para(doc, plan["goal"], size=11.2)
    add_callout(doc, "Сейчас", f"{items[plan['current_focus']]['id']} — {items[plan['current_focus']]['title']}. Параллельно наблюдаем стабильность {items[plan['parallel_validation']]['id']} — {items[plan['parallel_validation']]['title']}.")
    add_para(doc, "Важно: план не считает наличие кода завершением. Итоговый статус появляется только после приёмки в клинике и подтверждённой стабильной работы.", size=10.2, color=RED, bold=True)

    doc.add_page_break()
    add_heading(doc, "1. Как пользоваться планом", 1)
    add_para(doc, "В разговоре достаточно написать «План». Текущий реестр перечитывается, после чего показываются активный фокус, параллельная проверка, завершённые и незавершённые пункты, а также ближайшее действие.")
    add_heading(doc, "Команды", 2)
    commands = [
        ("make plan", "показать весь актуальный план и анализ статусов"),
        ("make plan-word", "заново собрать этот Word-файл из реестра"),
        ("npm run plan -- P0", "показать только критический приоритет P0"),
        ("npm run plan -- P0.1", "показать один конкретный пункт"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    for i, title in enumerate(("Команда", "Результат")):
        shade(table.rows[0].cells[i], NAVY)
        set_cell_text(table.rows[0].cells[i], title, size=9.2, color=WHITE, bold=True)
    repeat_header(table.rows[0])
    for command, meaning in commands:
        row = table.add_row()
        keep_row(row)
        set_cell_text(row.cells[0], command, size=9, color=TEAL, bold=True)
        set_cell_text(row.cells[1], meaning, size=9)

    add_heading(doc, "Статусы без самообмана", 2)
    status_table = doc.add_table(rows=1, cols=3)
    status_table.autofit = False
    for i, title in enumerate(("Статус", "Что означает", "Можно считать закрытым?")):
        shade(status_table.rows[0].cells[i], NAVY)
        set_cell_text(status_table.rows[0].cells[i], title, size=8.8, color=WHITE, bold=True)
    repeat_header(status_table.rows[0])
    for status in plan["status_model"]:
        row = status_table.add_row()
        keep_row(row)
        fill, color = STATUS_COLORS[status["code"]]
        shade(row.cells[0], fill)
        set_cell_text(row.cells[0], status["label"], size=8.7, color=color, bold=True)
        set_cell_text(row.cells[1], status["meaning"], size=8.7)
        closed = "Да" if status["code"] == "STABLE" else "Нет"
        set_cell_text(row.cells[2], closed, size=8.7, color=("1F5A32" if closed == "Да" else MUTED), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()
    add_heading(doc, "2. Сводная дорожная карта", 1)
    add_para(doc, "Порядок основан на пользовательской ценности для клиники. Техническая оптимизация и внешние интеграции не опережают стационар, документы, каталог и повторяемое подключение владельцев.")
    add_para(doc, "В сводной таблице ближайшее действие показано сокращённо; полный текст сохранён в карточке каждого пункта.", size=9.3, color=MUTED, italic=True)
    overview = doc.add_table(rows=1, cols=4)
    overview.autofit = False
    for i, title in enumerate(("Пункт", "Инициатива", "Статус", "Ближайшее действие")):
        shade(overview.rows[0].cells[i], NAVY)
        set_cell_text(overview.rows[0].cells[i], title, size=8.4, color=WHITE, bold=True)
    repeat_header(overview.rows[0])
    for item in plan["items"]:
        row = overview.add_row()
        keep_row(row)
        set_cell_text(row.cells[0], item["id"], size=8.4, color=TEAL, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], item["title"], size=8.4, bold=True)
        fill, color = STATUS_COLORS[item["status"]]
        shade(row.cells[2], fill)
        set_cell_text(row.cells[2], status_labels[item["status"]], size=8.0, color=color, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[3], compact_summary_action(item["next_action"]), size=8.1)

    add_item_detail(doc, items["P0.1"], status_labels, page_break=True)
    add_item_detail(doc, items["P0.2"], status_labels, page_break=True)
    add_item_detail(doc, items["P0.3"], status_labels, page_break=True)
    add_item_detail(doc, items["P0.4"], status_labels, page_break=True)
    add_item_detail(doc, items["P0.5"], status_labels, page_break=True)
    add_item_detail(doc, items["P0.6"], status_labels, page_break=True)

    doc.add_page_break()
    add_heading(doc, "8. P1 — развитие после основы", 1)
    add_para(doc, "P1 запускается параллельно с закрытием эксплуатационного P0.6. Первым выбран помощник на памяти фраз; документный редактор и архив по-прежнему зависят от «Документов 2.0», а сайт — от принятого owner-gateway.")
    for item_id in ("P1.1", "P1.2", "P1.3", "P1.4"):
        item = items[item_id]
        add_heading(doc, f"{item['id']}. {item['title']}", 2)
        add_para(doc, item["outcome"], size=10)
        add_para(doc, "Следующее действие: " + item["next_action"], size=9.4, color=MUTED, italic=True)
        for criterion in item["acceptance_criteria"]:
            add_bullet(doc, criterion)

    add_heading(doc, "9. P2 — подключаемое расширение после клинической основы", 1)
    add_callout(doc, "Граница приоритета", "Внешние лаборатории, коммуникации, платёжно-кассовый шлюз, сервисные модули, мультиклиника, OCR, локальная языковая модель и оптимизация чанка сохраняются в плане, но не забирают время у клинических P0/P1.", fill="FFF5DA", accent=GOLD)
    for item in (entry for entry in plan["items"] if entry["priority"] == "P2"):
        add_heading(doc, f"{item['id']}. {item['title']}", 2)
        add_para(doc, item["current_state"], size=10)
        add_para(doc, "Условие возврата: " + item["next_action"], size=9.4, color=MUTED, italic=True)

    add_heading(doc, "10. Порядок ближайших решений", 1)
    sequence = [
        "Принять локально «Документы 2.0»: версия шаблона, неизменяемый документ, PDF, подпись, отправка и архив пациента.",
        "После подтверждения выпустить P0.2 только с backup и заменой API/Web; состояние PostgreSQL, MinIO и других сервисов не перезапускать.",
        "Начать P0.3: сформировать безопасные очереди нормализации каталога и нулевых цен без изменения рабочих данных.",
        "Подготовить изолированную приёмку накладной: партии, остаток, цена, отмена и повторное открытие с доказанным восстановлением.",
        "Провести врачебную приёмку готового многосуточного листа на реальном пациенте и зафиксировать результат.",
        "Закрыть мобильную приёмку P0.6 и наблюдать кабинет владельца до статуса STABLE.",
        "После каталога перейти к визуальному редактору документов, затем к помощнику врача и ускорению frontend.",
        "После стабилизации клинической основы провести обследование платёжно-кассового контура, груминга и зоогостиницы; подключать их адаптерами и модулями без дублирования CRM.",
    ]
    for index, step in enumerate(sequence, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.7)
        set_run(p.add_run(f"{index}. "), size=10.5, color=TEAL, bold=True)
        set_run(p.add_run(step), size=10.5, color=INK)

    add_heading(doc, "11. Как обновляется план", 1)
    add_para(doc, "После каждой существенной доработки обновляются статус, доказательства, ближайшее действие и дата проверки. Переход между уровнями выполняется только на основании конкретного результата:")
    for text in (
        "коммит или diff подтверждает «Код готов»;",
        "сборка, тесты и визуальная проверка подтверждают «Проверено локально»;",
        "сотрудник и рабочее устройство подтверждают «Принято в клинике»;",
        "согласованный период эксплуатации подтверждает «Работает стабильно».",
    ):
        add_bullet(doc, text)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / f"Plan_dorabotki_TemichevVet_{plan['updated_at']}.docx"
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    build()
