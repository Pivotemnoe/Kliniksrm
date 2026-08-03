#!/usr/bin/env python3
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BRAND_NAVY = "173A5E"
BRAND_TEAL = "1F7A83"
BRAND_BLUE = "2E74B5"
INK = "18324A"
MUTED = "637789"
LIGHT = "EEF4F7"
LIGHT_TEAL = "E7F3F3"
GRID = "CBD8E2"
WHITE = "FFFFFF"
RISK = "9B1C1C"
GOLD = "946200"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            width = widths_dxa[min(i, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def create_decimal_numbering_id(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "100")
    spacing.set(qn("w:line"), "264")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(fonts)
    lvl.extend([start, num_fmt, lvl_text, suff, p_pr, r_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)


def set_run_font(run, name="Arial", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BRAND_BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.extend([fonts, color, underline])
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\))")


def add_inline(paragraph, text: str, base_size=10.5, base_color=INK) -> None:
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, color=base_color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Menlo", size=max(8.5, base_size - 1), color=BRAND_TEAL)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "EDF3F6")
            run._element.get_or_add_rPr().append(shd)
        else:
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link:
                add_hyperlink(paragraph, link.group(1), link.group(2))
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size, color=base_color)


def clean_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def add_field(paragraph, instruction: str, display="1") -> None:
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    paragraph.add_run()._r.append(fld_char)
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    paragraph.add_run()._r.append(instr_text)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._r.append(separate)
    value_run = paragraph.add_run(display)
    set_run_font(value_run, size=8, color=MUTED, bold=True)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(end)


def add_static_toc(doc: Document, source: Path) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("Оглавление")
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    r = intro.add_run("Полный отчёт: от текущего состояния и дорожной карты до приоритетов на 90 дней.")
    set_run_font(r, size=10, color=MUTED)
    headings = []
    for line in source.read_text(encoding="utf-8").splitlines():
        match = re.match(r"^##\s+(.+)$", line)
        if match:
            headings.append(clean_inline(match.group(1)))
    for heading in headings:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.0
        number = re.match(r"^(\d+\.)\s*(.*)$", heading)
        if number:
            r = p.add_run(number.group(1) + " ")
            set_run_font(r, size=9.4, color=BRAND_TEAL, bold=True)
            r = p.add_run(number.group(2))
            set_run_font(r, size=9.4, color=INK)
        else:
            r = p.add_run(heading)
            set_run_font(r, size=9.4, color=INK)
    doc.add_page_break()


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BRAND_NAVY, 16, 8),
        ("Heading 2", 13, BRAND_BLUE, 12, 6),
        ("Heading 3", 11.5, BRAND_TEAL, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.left_indent = Inches(0.50)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.10

    caption = styles["Caption"]
    caption.font.name = "Arial"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)


def configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.78)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.90)
        section.right_margin = Inches(0.90)
        section.header_distance = Inches(0.38)
        section.footer_distance = Inches(0.38)


def add_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TEMICHEVVET CRM  /  ПОЛНЫЙ ПРОДУКТОВЫЙ АУДИТ")
    set_run_font(run, size=8.2, color=MUTED, bold=True)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), BRAND_TEAL)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("3 августа 2026  •  стр. ")
    set_run_font(run, size=8, color=MUTED)
    add_field(p, "PAGE")


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("ПРОДУКТОВЫЙ АУДИТ  /  2026")
    set_run_font(r, size=10, color=BRAND_TEAL, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("TemichevVet CRM")
    set_run_font(r, size=31, color=BRAND_NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("Полный аудит функций, дорожной карты, UX, документов, стационара и продуктовых разрывов")
    set_run_font(r, size=15, color=MUTED)

    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2340, 2340, 2340, 2340])
    set_repeat_table_header(table.rows[0])
    metrics = [
        ("85–86%", "широта относительно Vet.AF"),
        ("88–91%", "готовность одной клиники"),
        ("72%", "документы и шаблоны"),
        ("60%", "стационарный лист"),
    ]
    for i, (value, label) in enumerate(metrics):
        cell = table.cell(0, i)
        set_cell_shading(cell, BRAND_NAVY if i == 0 else LIGHT_TEAL)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(value)
        set_run_font(r, size=15, color=WHITE if i == 0 else BRAND_NAVY, bold=True)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=8.2, color=WHITE if i == 0 else MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Главный вывод")
    set_run_font(r, size=12, color=BRAND_TEAL, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.20
    r = p.add_run("CRM уже является большим рабочим продуктом. Следующий прирост ценности дают не новые разделы, а углубление четырёх ключевых контуров: стационар, документы, качество каталога и путь владельца животного.")
    set_run_font(r, size=13, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Проверенный срез")
    set_run_font(r, size=9, color=MUTED, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("efc6c95 · origin/main и текущий HEAD · 3 августа 2026 года")
    set_run_font(r, size=10.5, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Без изменения клинических записей, Docker volumes, рабочего сервера, TECNO и флешки")
    set_run_font(r, size=9.2, color=MUTED, italic=True)
    doc.add_page_break()


def parse_blocks(lines: list[str]):
    i = 0
    paragraph_lines: list[str] = []

    def flush_para():
        nonlocal paragraph_lines
        if paragraph_lines:
            joined = " ".join(s.strip().rstrip("  ") for s in paragraph_lines).strip()
            paragraph_lines = []
            if joined:
                return ("paragraph", joined)
        return None

    while i < len(lines):
        line = lines[i].rstrip("\n")
        if not line.strip():
            block = flush_para()
            if block:
                yield block
            i += 1
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            block = flush_para()
            if block:
                yield block
            yield ("heading", len(heading.group(1)), heading.group(2).strip())
            i += 1
            continue
        image = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", line.strip())
        if image:
            block = flush_para()
            if block:
                yield block
            yield ("image", image.group(1), image.group(2))
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|(?:\s*:?-{3,}:?\s*\|)+\s*$", lines[i + 1].strip()):
            block = flush_para()
            if block:
                yield block
            table_lines = [line, lines[i + 1].rstrip("\n")]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip("\n"))
                i += 1
            rows = []
            for idx, row in enumerate(table_lines):
                if idx == 1:
                    continue
                cells = [c.strip() for c in row.strip().strip("|").split("|")]
                rows.append(cells)
            yield ("table", rows)
            continue
        if line.lstrip().startswith(">"):
            block = flush_para()
            if block:
                yield block
            quote_lines = []
            while i < len(lines) and lines[i].lstrip().startswith(">"):
                quote_lines.append(lines[i].lstrip()[1:].strip())
                i += 1
            yield ("quote", " ".join(quote_lines))
            continue
        list_match = re.match(r"^\s*([-*])\s+(.+)$", line)
        num_match = re.match(r"^\s*(\d+)\.\s+(.+)$", line)
        if list_match or num_match:
            block = flush_para()
            if block:
                yield block
            kind = "bullet" if list_match else "number"
            text = (list_match or num_match).group(2)
            yield ("list", kind, text)
            i += 1
            continue
        paragraph_lines.append(line)
        i += 1
    block = flush_para()
    if block:
        yield block


def width_distribution(rows: list[list[str]]) -> list[int]:
    n = max(len(r) for r in rows)
    if n == 2:
        return [2800, 6560]
    if n == 3:
        lengths = [max(len(clean_inline(r[i])) if i < len(r) else 0 for r in rows) for i in range(n)]
        if lengths[1] < 18:
            return [3000, 1600, 4760]
        return [2500, 2500, 4360]
    if n == 4:
        return [2050, 1250, 1900, 4160]
    return [CONTENT_DXA // n] * (n - 1) + [CONTENT_DXA - (CONTENT_DXA // n) * (n - 1)]


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    widths = width_distribution(rows)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx in range(ncols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            value = row[c_idx] if c_idx < len(row) else ""
            add_inline(p, value, base_size=8.3 if ncols >= 4 else 8.8, base_color=WHITE if r_idx == 0 else INK)
            if r_idx == 0:
                set_cell_shading(cell, BRAND_NAVY)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "F7FAFB")
            if c_idx == 1 and re.fullmatch(r"\s*\d+(?:[–-]\d+)?%?\s*", clean_inline(value)):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.15
    add_inline(p, text, base_size=10, base_color=INK)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_TEAL)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BRAND_TEAL)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def add_content(doc: Document, source: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    first_h1_seen = False
    content_started = False
    first_main_section = True
    current_num_id = None
    for block in parse_blocks(lines):
        kind = block[0]
        if not (kind == "list" and block[1] == "number"):
            current_num_id = None
        if kind == "heading":
            _, level, text = block
            if level == 1:
                if not first_h1_seen:
                    first_h1_seen = True
                    continue
                level = 1
            if level == 2:
                content_started = True
                if not first_main_section:
                    doc.add_page_break()
                first_main_section = False
                style = "Heading 1"
            elif level == 3:
                style = "Heading 2"
            else:
                style = "Heading 3"
            p = doc.add_paragraph(style=style)
            add_inline(p, text, base_size={"Heading 1": 16, "Heading 2": 13, "Heading 3": 11.5}[style], base_color={"Heading 1": BRAND_NAVY, "Heading 2": BRAND_BLUE, "Heading 3": BRAND_TEAL}[style])
            for run in p.runs:
                run.bold = True
        elif not content_started:
            continue
        elif kind == "paragraph":
            p = doc.add_paragraph()
            add_inline(p, block[1])
        elif kind == "list":
            _, list_kind, text = block
            p = doc.add_paragraph(style="List Bullet" if list_kind == "bullet" else "List Number")
            if list_kind == "number":
                if current_num_id is None:
                    current_num_id = create_decimal_numbering_id(doc)
                apply_numbering(p, current_num_id)
            add_inline(p, text)
        elif kind == "quote":
            add_quote(doc, block[1])
        elif kind == "table":
            add_markdown_table(doc, block[1])
        elif kind == "image":
            _, alt, rel_path = block
            image_path = source.parent / rel_path
            if image_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.keep_with_next = True
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(3)
                shape = p.add_run().add_picture(str(image_path), width=Inches(6.45))
                shape._inline.docPr.set("descr", alt or image_path.stem)
                shape._inline.docPr.set("title", alt or image_path.stem)
                caption = doc.add_paragraph(style="Caption")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.keep_together = True
                caption.paragraph_format.space_after = Pt(10)
                caption.add_run(alt)


def add_update_fields_flag(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def build(source: Path, target: Path) -> None:
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    for section in doc.sections:
        add_header_footer(section)
    add_cover(doc)
    add_static_toc(doc, source)
    add_content(doc, source)
    add_update_fields_flag(doc)
    core = doc.core_properties
    core.title = "TemichevVet CRM — полный продуктовый аудит"
    core.subject = "Функции, дорожная карта, UX, документы, стационар и продуктовые разрывы"
    core.author = "TemichevVet"
    core.keywords = "TemichevVet, CRM, продуктовый аудит, Vet.AF, стационар, документы"
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: build_full_product_audit_docx.py SOURCE.md TARGET.docx")
    build(Path(sys.argv[1]).resolve(), Path(sys.argv[2]).resolve())
